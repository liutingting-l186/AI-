"""
文档预处理：PDF/Word/文本 → OCR（扫描件）→ 细粒度切分 → Embedding → SQLite入库
支持：.pdf（文本型/扫描件）/ .txt / .md / .doc / .docx
"""
import os
import sqlite3
import numpy as np
import re
from pathlib import Path
from sentence_transformers import SentenceTransformer
from typing import List, Dict, Optional
import pdfplumber
import fitz  # PyMuPDF
import torch  # 移至顶部确保全局可用

# ======== 添加在文件最顶部（所有import之前）========
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'  # 跳过模型源检查
os.environ['PADDLEOCR_LOG_LEVEL'] = 'ERROR'  # 仅显示错误日志
# ======== 以上代码必须放在最顶部 ========

# 配置
DB_PATH = "./knowledge.db"
DOC_DIR = "./documents"
TEXT_OUTPUT_DIR = "./extracted_text"  # 新增：文本提取物保存目录
EMBEDDING_MODEL_PATH = "./models/bge-large-zh-v1.5"
USE_OCR = True  # 是否启用OCR（扫描件PDF必需）
OCR_ENGINE = 'easyocr'  # 使用EasyOCR替代PaddleOCR（兼容性更好）

# 尝试导入Word处理依赖（按需加载）
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ 未安装python-docx，.docx文件处理将不可用 (pip install python-docx)")

try:
    import textract
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False
    print("⚠️ 未安装textract，.doc文件处理将不可用 (pip install textract)")

class DocumentProcessor:
    def __init__(
        self,
        doc_dir: str = DOC_DIR,
        db_path: str = DB_PATH,
        text_output_dir: str = TEXT_OUTPUT_DIR,
        embedding_model_path: str = EMBEDDING_MODEL_PATH,
    ):
        self.doc_dir = Path(doc_dir)
        self.db_path = db_path
        self.text_output_dir = Path(text_output_dir)
        self.text_output_dir.mkdir(parents=True, exist_ok=True)
        
        self.conn = sqlite3.connect(self.db_path)
        self._init_db()
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model = SentenceTransformer(embedding_model_path, device=self.device)
        print(f"✓ Embedding模型加载完成 (device: {self.device})")
        
        # OCR懒加载（仅当需要时初始化）
        self.ocr_engine = None
    
    def _init_db(self):
        """创建SQLite表结构"""
        self.conn.execute('''
            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                doc_id TEXT NOT NULL,
                doc_name TEXT NOT NULL,
                chunk_text TEXT NOT NULL,
                page_num INTEGER DEFAULT -1,
                section TEXT DEFAULT '',
                clause TEXT DEFAULT '',
                business_domain TEXT DEFAULT '通用',
                permission_tag TEXT DEFAULT 'public',
                embedding BLOB NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_permission ON chunks(permission_tag)')
        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_domain ON chunks(business_domain)')
        self.conn.execute('CREATE INDEX IF NOT EXISTS idx_doc_page ON chunks(doc_id, page_num)')
        self.conn.commit()
        print(f"✓ SQLite数据库初始化完成: {self.db_path}")
    
    def _infer_domain(self, filename: str) -> str:
        filename_lower = filename.lower()
        if any(kw in filename_lower for kw in ['危化', '危险', '剧毒', '爆炸']):
            return "危化品运输"
        elif any(kw in filename_lower for kw in ['客运', '班车', '旅游包车']):
            return "道路客运"
        elif any(kw in filename_lower for kw in ['货运', '物流', '货车']):
            return "道路货运"
        elif any(kw in filename_lower for kw in ['事故', '案例', '复盘']):
            return "事故案例"
        return "通用监管"
    
    def _extract_clause_number(self, text: str) -> str:
        patterns = [
            r'第[零一二三四五六七八九十百千]+条',
            r'第\d+条',
            r'第[零一二三四五六七八九十百千]+章',
            r'第\d+章'
        ]
        for pattern in patterns:
            match = re.search(pattern, text[:100])
            if match:
                return match.group(0)
        return ""
    
    def _split_chunks_with_page(self, pages: List[Dict]) -> List[Dict]:
        """智能切分（避免跨页断裂）"""
        chunks = []
        current_chunk = []
        current_page = -1
        char_count = 0
        
        for page_data in pages:
            page_num = page_data["page_num"]
            lines = [line.strip() for line in page_data["text"].split('\n') if line.strip()]
            
            for line in lines:
                # 跳过页眉页脚
                if len(line) < 20 and re.search(r'第\s*\d+\s*页|page\s*\d+', line, re.I):
                    continue
                
                # 检测新条款开始（触发强制切分）
                if re.match(r'^第[零一二三四五六七八九十百千\d]+[条章]', line):
                    if current_chunk:
                        chunk_text = '\n'.join(current_chunk)
                        chunks.append({
                            "chunk_id": f"{page_data['doc_id']}_{len(chunks)}",
                            "doc_id": page_data['doc_id'],
                            "doc_name": page_data['doc_name'],
                            "chunk_text": chunk_text,
                            "page_num": current_page,
                            "section": self._extract_clause_number(chunk_text),
                            "clause": self._extract_clause_number(chunk_text),
                            "business_domain": page_data['business_domain'],
                            "permission_tag": page_data['permission_tag']
                        })
                        current_chunk = []
                        char_count = 0
                
                current_chunk.append(line)
                char_count += len(line)
                current_page = page_num
                
                if char_count > 800 and not re.match(r'^第[零一二三四五六七八九十百千\d]+[条章]', line):
                    chunk_text = '\n'.join(current_chunk)
                    chunks.append({
                        "chunk_id": f"{page_data['doc_id']}_{len(chunks)}",
                        "doc_id": page_data['doc_id'],
                        "doc_name": page_data['doc_name'],
                        "chunk_text": chunk_text,
                        "page_num": current_page,
                        "section": self._extract_clause_number(chunk_text),
                        "clause": self._extract_clause_number(chunk_text),
                        "business_domain": page_data['business_domain'],
                        "permission_tag": page_data['permission_tag']
                    })
                    current_chunk = []
                    char_count = 0
        
        if current_chunk:
            chunk_text = '\n'.join(current_chunk)
            chunks.append({
                "chunk_id": f"{pages[0]['doc_id']}_{len(chunks)}",
                "doc_id": pages[0]['doc_id'],
                "doc_name": pages[0]['doc_name'],
                "chunk_text": chunk_text,
                "page_num": current_page,
                "section": self._extract_clause_number(chunk_text),
                "clause": self._extract_clause_number(chunk_text),
                "business_domain": pages[0]['business_domain'],
                "permission_tag": pages[0]['permission_tag']
            })
        
        chunks = [c for c in chunks if len(c["chunk_text"]) > 50]
        return chunks
    
    def _parse_pdf_with_pdfplumber(self, filepath: str) -> Optional[List[Dict]]:
        """主方案：pdfplumber（文本型PDF）"""
        try:
            pages = []
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text(
                        keep_blank_chars=True,
                        layout=True,
                        x_tolerance=1,
                        y_tolerance=1
                    ) or ""
                    
                    # 表格转Markdown
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table and len(table) > 1:
                                md_table = "\n\n表格内容：\n"
                                for row in table[:3]:
                                    md_table += " | ".join([str(cell) if cell else "" for cell in row[:4]]) + "\n"
                                text += md_table
                    
                    pages.append({"page_num": page_num, "text": text.strip()})
            return pages if any(p["text"].strip() for p in pages) else None
        except Exception as e:
            print(f"  ⚠️ pdfplumber解析失败: {str(e)[:80]}")
            return None
    
    def _parse_pdf_with_fitz(self, filepath: str) -> Optional[List[Dict]]:
        """降级方案：PyMuPDF（文本型PDF备选）"""
        try:
            pages = []
            doc = fitz.open(filepath)
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
                pages.append({"page_num": page_num, "text": text.strip()})
            doc.close()
            return pages if any(p["text"].strip() for p in pages) else None
        except Exception as e:
            print(f"  ⚠️ PyMuPDF解析失败: {str(e)[:80]}")
            return None
    
    def _init_ocr(self):
        """懒加载OCR引擎（使用EasyOCR，自动检测GPU）"""
        if self.ocr_engine is not None:
            return self.ocr_engine
        
        print("  → 首次使用OCR，加载EasyOCR模型（约20-40秒）...")
        try:
            import easyocr
            # 根据GPU可用性自动配置
            use_gpu = torch.cuda.is_available()
            print(f"  → OCR设备选择: {'GPU' if use_gpu else 'CPU'}")
            # EasyOCR创建识别器（自动下载中文模型）
            self.ocr_engine = easyocr.Reader(['ch_sim'], gpu=use_gpu)  # 简体中文识别
            
            print("  ✓ OCR引擎加载完成")
            return self.ocr_engine
        except ImportError:
            raise ImportError(
                "❌ 未安装EasyOCR，请先安装:\n"
                "   pip install easyocr")
        except Exception as e:
            raise RuntimeError(f"OCR初始化失败: {str(e)}")
    
    def _parse_pdf_with_ocr(self, filepath: str) -> Optional[List[Dict]]:
        """OCR方案：处理扫描件PDF"""
        if not USE_OCR:
            print("  ⚠️ OCR功能未启用 (USE_OCR=False)")
            return None
        
        try:
            # 检查pdf2image依赖
            try:
                import pdf2image
            except ImportError:
                raise ImportError(
                    "❌ 未安装pdf2image，请先安装:\n"
                    "   pip install pdf2image -i https://mirrors.aliyun.com/pypi/simple/   --trusted-host mirrors.aliyun.com\n"
                    "   Windows用户还需安装poppler: https://github.com/oschwartz10612/poppler-windows/releases/  "
                )
            
            # Windows特殊处理：自动检测poppler路径
            poppler_path = None
            if os.name == 'nt':  # Windows
                # 尝试常见路径
                common_paths = [
                    r"C:\Program Files\poppler\Library\bin",
                    r"C:\poppler\Library\bin",
                    r"D:\poppler\Library\bin",
                    # 从环境变量获取
                    os.environ.get("POPPLER_PATH")
                ]
                for path in common_paths:
                    if path and os.path.exists(path) and "pdftoppm.exe" in os.listdir(path):
                        poppler_path = path
                        print(f"  → 自动检测到poppler路径: {poppler_path}")
                        break
                if not poppler_path:
                    print("  ⚠️ 未找到poppler路径，尝试使用系统PATH（如失败请手动安装poppler）")
            
            # PDF转图片
            print(f"  → OCR处理中（PDF转图片）...")
            images = pdf2image.convert_from_path(
                filepath,
                dpi=400,  # 200 DPI平衡速度与精度
                poppler_path=poppler_path if poppler_path else None
            )
            
            # 初始化OCR
            ocr = self._init_ocr()
            
            # 逐页OCR
            pages = []
            for page_num, image in enumerate(images, 1):
                print(f"    OCR识别 P{page_num}/{len(images)}...", end="\r")
                # EasyOCR 直接接收 PIL Image 或 numpy 数组
                image_array = np.array(image)
                results = ocr.readtext(image_array, text_threshold=0.5)  # text_threshold 筛选置信度
                
                if not results:
                    text = ""
                else:
                    # EasyOCR 返回格式：[(bbox, text, confidence), ...]
                    text_lines = [result[1] for result in results if result[1].strip()]
                    text = "\n".join(text_lines)
                pages.append({"page_num": page_num, "text": text.strip()})
            
            print(f"\n  ✓ OCR完成 {len(pages)} 页")
            return pages if any(p["text"].strip() for p in pages) else None
            
        except Exception as e:
            error_msg = str(e)[:100]
            print(f"  ⚠️ OCR解析失败: {error_msg}")
            # 仅在调试时显示详细堆栈
            # import traceback; traceback.print_exc()
            return None
    
    def _parse_pdf(self, filepath: str) -> List[Dict]:
        """PDF解析主入口（三重降级策略）"""
        print(f"  尝试解析PDF...")
        
        # 策略1: pdfplumber（文本型PDF）
        pages = self._parse_pdf_with_pdfplumber(filepath)
        if pages:
            print(f"  ✓ pdfplumber成功: {len(pages)} 页 (文本型PDF)")
            return pages
        
        # 策略2: PyMuPDF（文本型PDF备选）
        print(f"  → 回退到PyMuPDF...")
        pages = self._parse_pdf_with_fitz(filepath)
        if pages:
            print(f"  ✓ PyMuPDF成功: {len(pages)} 页 (文本型PDF)")
            return pages
        
        # 策略3: OCR（扫描件PDF）
        if USE_OCR:
            print(f"  → 回退到OCR（扫描件PDF）...")
            pages = self._parse_pdf_with_ocr(filepath)
            if pages:
                print(f"  ✓ OCR成功: {len(pages)} 页 (扫描件PDF)")
                return pages
        
        raise ValueError("PDF解析失败：文本型与OCR方案均未提取到有效文本")
    
    def _parse_text(self, filepath: str) -> List[Dict]:
        """文本文件解析"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(filepath, 'r', encoding='gbk') as f:
                text = f.read()
        
        page_size = 2000
        pages = []
        for i in range(0, len(text), page_size):
            pages.append({
                "page_num": i // page_size + 1,
                "text": text[i:i+page_size].strip()
            })
        return pages
    
    def _parse_docx(self, filepath: str) -> Optional[List[Dict]]:
        """Word .docx 文件解析"""
        if not HAS_DOCX:
            print("  ⚠️ 未安装python-docx，跳过.docx文件")
            return None
        
        try:
            doc = Document(filepath)
            # 提取所有段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            # 提取表格内容（简化处理）
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        full_text += "\n\n表格: " + " | ".join(cells)
            
            if not full_text.strip():
                print("  ⚠️ .docx文件内容为空")
                return None
            
            # 整个文档视为单页（Word无原生分页概念）
            return [{"page_num": 1, "text": full_text.strip()}]
        except Exception as e:
            print(f"  ⚠️ .docx解析失败: {str(e)[:80]}")
            return None
    
    def _parse_doc(self, filepath: str) -> Optional[List[Dict]]:
        """Word .doc 文件解析（通过textract）"""
        if not HAS_TEXTRACT:
            print("  ⚠️ 未安装textract，无法处理.doc文件（建议转换为.docx）")
            return None
        
        try:
            # textract自动处理.doc格式
            text = textract.process(str(filepath), encoding='utf-8')
            decoded_text = text.decode('utf-8').strip()
            
            if not decoded_text:
                print("  ⚠️ .doc文件内容为空")
                return None
            
            # 整个文档视为单页
            return [{"page_num": 1, "text": decoded_text}]
        except Exception as e:
            print(f"  ⚠️ .doc解析失败: {str(e)[:80]}")
            return None
    
    def _save_extracted_text(self, filepath: Path, pages: List[Dict]):
        """保存提取的原始文本到本地（用于校对）- 健壮版"""
        try:
            # 合并所有页面文本（保留分页标记）
            full_text = "\n\n--- Page Break ---\n\n".join(
                [f"Page {p['page_num']}:\n{p['text']}" for p in pages]
            )
            
            # === 健壮的路径处理：规范化 + 安全回退 ===
            try:
                # 1. 规范化路径（解析符号链接、统一大小写/斜杠）
                doc_dir_abs = self.doc_dir.resolve()
                file_abs = Path(filepath).resolve()
                
                # 2. 尝试获取相对路径（捕获所有异常）
                try:
                    rel_path = file_abs.relative_to(doc_dir_abs)
                    # 3. 扁平化保存：将目录结构转换为文件名（避免深层嵌套）
                    #    例如: 子目录/文件.pdf -> extracted_text/子目录_文件.txt
                    safe_name = str(rel_path.parent).replace(os.sep, '_').replace('/', '_')
                    if safe_name and safe_name != '.':
                        output_name = f"{safe_name}_{file_abs.stem}.txt"
                    else:
                        output_name = f"{file_abs.stem}.txt"
                except (ValueError, AttributeError):
                    # 回退：仅用文件名（不同目录同名文件可能覆盖，但概率低）
                    output_name = f"{file_abs.stem}.txt"
            except Exception:
                # 极端情况：直接用原始文件名
                output_name = f"{filepath.stem}.txt"
            
            # 4. 清理Windows非法字符（避免保存失败）
            illegal_chars = r'<>:"/\|?*'
            for char in illegal_chars:
                output_name = output_name.replace(char, '_')
            
            # 5. 生成最终输出路径（始终在text_output_dir下平铺保存）
            output_path = self.text_output_dir / output_name
            
            # 6. 自动创建目录（parents=True 确保父目录存在）
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 7. 保存文本
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            
            print(f"  ✓ 原始文本已保存: {output_path.name}")
        except Exception as e:
            print(f"  ⚠️ 文本保存失败: {str(e)[:100]}")
            
    def _file_to_chunks(self, filepath: str) -> List[Dict]:
        """统一入口：解析并切分"""
        filepath = Path(filepath)
        doc_id = filepath.stem
        doc_name = filepath.name
        business_domain = self._infer_domain(doc_name)
        permission_tag = "public"
        
        # 解析内容（根据文件类型分发）
        suffix = filepath.suffix.lower()
        if suffix == '.pdf':
            pages = self._parse_pdf(str(filepath))
        elif suffix in ['.txt', '.md']:
            pages = self._parse_text(str(filepath))
        elif suffix == '.docx':
            pages = self._parse_docx(str(filepath))
        elif suffix == '.doc':
            pages = self._parse_doc(str(filepath))
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")
        
        if not pages:
            raise ValueError("未提取到有效文本内容")
        
        # 保存原始文本用于校对
        self._save_extracted_text(filepath, pages)
        
        # 添加元数据
        page_data = [{
            "page_num": p["page_num"],
            "text": p["text"],
            "doc_id": doc_id,
            "doc_name": doc_name,
            "business_domain": business_domain,
            "permission_tag": permission_tag
        } for p in pages]
        
        # 智能切分
        chunks = self._split_chunks_with_page(page_data)
        print(f"  → 切分为 {len(chunks)} 个chunks (平均{sum(len(c['chunk_text']) for c in chunks)//len(chunks) if chunks else 0}字符)")
        return chunks
    
    def process_documents(self, doc_dir: Optional[str] = None):
        """批量处理文档（递归遍历子目录）"""
        doc_dir = Path(doc_dir) if doc_dir else self.doc_dir
        if not doc_dir.exists():
            doc_dir.mkdir(parents=True, exist_ok=True)
            print(f"⚠️ 文档目录不存在，已创建: {doc_dir}")
            return
        
        # 递归查找所有支持的文件（跳过隐藏文件）
        supported_exts = {'.pdf', '.txt', '.md', '.doc', '.docx'}
        doc_files = [
            f for f in doc_dir.rglob("*") 
            if f.is_file() 
            and f.suffix.lower() in supported_exts
            and not f.name.startswith('.')
        ]
        
        if not doc_files:
            print(f"⚠️ 未找到文档（支持格式: {', '.join(supported_exts)}）")
            print(f"   请将文档放入 {doc_dir} 及其子目录")
            return
        
        print(f"✓ 找到 {len(doc_files)} 个文档（含子目录），开始处理...")
        all_chunks = []
        
        for filepath in sorted(doc_files):  # 按路径排序便于追踪
            print(f"\n📄 处理: {filepath.relative_to(doc_dir)}")
            try:
                chunks = self._file_to_chunks(str(filepath))
                all_chunks.extend(chunks)
                print(f"  ✓ 成功: {len(chunks)} chunks")
            except Exception as e:
                print(f"  ✗ 失败: {str(e)}")
                continue
        
        if not all_chunks:
            print("✗ 未生成任何chunks")
            return
        
        # 生成Embedding（自动使用GPU）
        print(f"\n✓ 生成 {len(all_chunks)} 个chunks的Embedding (device: {self.device})...")
        texts = [c["chunk_text"] for c in all_chunks]
        embeddings = self.model.encode(
            texts, 
            batch_size=32, 
            show_progress_bar=True,
            normalize_embeddings=True,
            device=self.device  # 显式指定设备
        )
        
        # 写入SQLite
        print(f"✓ 写入SQLite数据库...")
        cursor = self.conn.cursor()
        for i, (chunk, emb) in enumerate(zip(all_chunks, embeddings)):
            emb_bytes = emb.astype(np.float32).tobytes()
            cursor.execute('''
                INSERT OR REPLACE INTO chunks 
                (chunk_id, doc_id, doc_name, chunk_text, page_num, section, clause,
                 business_domain, permission_tag, embedding)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                chunk["chunk_id"], chunk["doc_id"], chunk["doc_name"], 
                chunk["chunk_text"], chunk["page_num"], chunk["section"], chunk["clause"],
                chunk["business_domain"], chunk["permission_tag"], emb_bytes
            ))
            if (i + 1) % 100 == 0:
                self.conn.commit()
                print(f"  → 已写入 {i+1}/{len(all_chunks)} chunks")
        self.conn.commit()
        print(f"\n🎉 全部完成！共入库 {len(all_chunks)} 个chunks")
        print(f"   数据库: {Path(self.db_path).absolute()}")
        print(f"   原始文本备份: {self.text_output_dir.absolute()}")
    
    def close(self):
        self.conn.close()

# ============ 主程序 ============
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("-i", "--input-dir", default=DOC_DIR)
    parser.add_argument("-o", "--db-path", default=DB_PATH)
    parser.add_argument("--text-output-dir", default=TEXT_OUTPUT_DIR)
    args = parser.parse_args()

    doc_dir = Path(args.input_dir)
    supported_exts = {'.pdf', '.txt', '.md', '.doc', '.docx'}
    has_files = any(
        f.is_file() and f.suffix.lower() in supported_exts and not f.name.startswith('.')
        for f in doc_dir.rglob("*")
    )
    
    if not has_files:
        print(f"⚠️ 警告: 文档目录 '{DOC_DIR}' 及其子目录中未找到支持的文档")
        print(f"   支持格式: {', '.join(supported_exts)}")
        print("   请将文档放入该目录后重新运行")
        exit(1)
    else:
        # 统计各类文件数量
        ext_counts = {}
        for f in doc_dir.rglob("*"):
            if f.is_file() and f.suffix.lower() in supported_exts and not f.name.startswith('.'):
                ext = f.suffix.lower()
                ext_counts[ext] = ext_counts.get(ext, 0) + 1
        
        print(f"✓ 检测到文档: {sum(ext_counts.values())} 个（含子目录）")
        for ext, count in sorted(ext_counts.items()):
            print(f"  - {ext}: {count} 个")
    
    # 执行处理
    processor = DocumentProcessor(
        doc_dir=str(doc_dir),
        db_path=args.db_path,
        text_output_dir=args.text_output_dir,
    )
    processor.process_documents(str(doc_dir))
    processor.close()
    
    # 验证结果
    print("\n✓ 验证入库结果:")
    conn = sqlite3.connect(args.db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM chunks")
    total = cursor.fetchone()[0]
    cursor.execute("SELECT doc_name, page_num, section, clause, chunk_text FROM chunks LIMIT 3")
    samples = cursor.fetchall()
    conn.close()
    
    print(f"  总chunks数: {total}")
    for i, (doc, page, section, clause, text) in enumerate(samples, 1):
        print(f"\n  [{i}] {doc} P{page} | 章节:{section} 条款:{clause}")
        print(f"      {text[:100]}...")
