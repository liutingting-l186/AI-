"""
阶段1：文档文本提取器
功能：从PDF/Word/文本中提取原始内容，保存为JSON格式供人工校对
输出：./extracted_text/{doc_id}.json（含元数据+分页文本）
特点：
  ✅ 完整保留原有解析逻辑（PDF/OCR/Word三重降级）
  ✅ 不生成Embedding、不写入数据库
  ✅ JSON格式便于人工校对（直接编辑text字段）
  ✅ 保留所有元数据（doc_id/page_num/domain等）
"""
import os
import json
import numpy as np
import re
from pathlib import Path
from typing import List, Dict, Optional
import pdfplumber
import fitz  # PyMuPDF
import torch

# ======== 环境变量（必须在最顶部）========
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'
os.environ['PADDLEOCR_LOG_LEVEL'] = 'ERROR'
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"  # 强制CPU避免驱动问题
# ======== 以上必须在import torch之前 ========

# 配置
DOC_DIR = "./documents"
TEXT_OUTPUT_DIR = "./extracted_text"  # 校对文本保存目录
USE_OCR = True  # 是否启用OCR（扫描件PDF必需）

# 尝试导入Word处理依赖（按需加载）
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ 未安装python-docx，.docx文件处理将不可用 (pip install python-docx)")

try:
    import textract
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False
    print("⚠️ 未安装textract，.doc文件处理将不可用 (pip install textract)")


class TextExtractor:
    def __init__(self):
        self.text_output_dir = Path(TEXT_OUTPUT_DIR)
        self.text_output_dir.mkdir(parents=True, exist_ok=True)
        self.ocr_engine = None  # OCR懒加载
    
    def _infer_domain(self, filename: str) -> str:
        filename_lower = filename.lower()
        if any(kw in filename_lower for kw in ['危化', '危险', '剧毒', '爆炸']):
            return "危化品运输"
        elif any(kw in filename_lower for kw in ['客运', '班车', '旅游包车']):
            return "道路客运"
        elif any(kw in filename_lower for kw in ['货运', '物流', '货车']):
            return "道路货运"
        elif any(kw in filename_lower for kw in ['事故', '案例', '复盘']):
            return "事故案例"
        return "通用监管"
    
    def _parse_pdf_with_pdfplumber(self, filepath: str) -> Optional[List[Dict]]:
        """主方案：pdfplumber（文本型PDF）"""
        try:
            pages = []
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text(
                        keep_blank_chars=True,
                        layout=True,
                        x_tolerance=1,
                        y_tolerance=1
                    ) or ""
                    
                    # 表格转Markdown
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table and len(table) > 1:
                                md_table = "\n\n表格内容：\n"
                                for row in table[:3]:
                                    md_table += " | ".join([str(cell) if cell else "" for cell in row[:4]]) + "\n"
                                text += md_table
                    
                    pages.append({"page_num": page_num, "text": text.strip()})
            return pages if any(p["text"].strip() for p in pages) else None
        except Exception as e:
            print(f"  ⚠️ pdfplumber解析失败: {str(e)[:80]}")
            return None
    
    def _parse_pdf_with_fitz(self, filepath: str) -> Optional[List[Dict]]:
        """降级方案：PyMuPDF（文本型PDF备选）"""
        try:
            pages = []
            doc = fitz.open(filepath)
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
                pages.append({"page_num": page_num, "text": text.strip()})
            doc.close()
            return pages if any(p["text"].strip() for p in pages) else None
        except Exception as e:
            print(f"  ⚠️ PyMuPDF解析失败: {str(e)[:80]}")
            return None
    
    def _init_ocr(self):
        """懒加载OCR引擎（使用EasyOCR）"""
        if self.ocr_engine is not None:
            return self.ocr_engine
        
        print("  → 首次使用OCR，加载EasyOCR模型（约20-40秒）...")
        try:
            import easyocr
            use_gpu = torch.cuda.is_available()
            print(f"  → OCR设备选择: {'GPU' if use_gpu else 'CPU'}")
            self.ocr_engine = easyocr.Reader(['ch_sim'], gpu=use_gpu)
            print("  ✓ OCR引擎加载完成")
            return self.ocr_engine
        except ImportError:
            raise ImportError("❌ 未安装EasyOCR，请先安装: pip install easyocr")
        except Exception as e:
            raise RuntimeError(f"OCR初始化失败: {str(e)}")
    
    def _parse_pdf_with_ocr(self, filepath: str) -> Optional[List[Dict]]:
        """OCR方案：处理扫描件PDF"""
        if not USE_OCR:
            print("  ⚠️ OCR功能未启用 (USE_OCR=False)")
            return None
        
        try:
            try:
                import pdf2image
            except ImportError:
                raise ImportError(
                    "❌ 未安装pdf2image，请先安装:\n"
                    "   pip install pdf2image -i https://mirrors.aliyun.com/pypi/simple/ --trusted-host mirrors.aliyun.com\n"
                    "   Windows用户还需安装poppler: https://github.com/oschwartz10612/poppler-windows/releases/"
                )
            
            # Windows poppler路径检测
            poppler_path = None
            if os.name == 'nt':
                common_paths = [
                    r"C:\Program Files\poppler\Library\bin",
                    r"C:\poppler\Library\bin",
                    r"D:\poppler\Library\bin",
                    os.environ.get("POPPLER_PATH")
                ]
                for path in common_paths:
                    if path and os.path.exists(path) and "pdftoppm.exe" in os.listdir(path):
                        poppler_path = path
                        print(f"  → 自动检测到poppler路径: {poppler_path}")
                        break
                if not poppler_path:
                    print("  ⚠️ 未找到poppler路径，尝试使用系统PATH")
            
            # PDF转图片
            print(f"  → OCR处理中（PDF转图片）...")
            images = pdf2image.convert_from_path(
                filepath,
                dpi=400,
                poppler_path=poppler_path if poppler_path else None
            )
            
            # 逐页OCR
            ocr = self._init_ocr()
            pages = []
            for page_num, image in enumerate(images, 1):
                print(f"    OCR识别 P{page_num}/{len(images)}...", end="\r")
                image_array = np.array(image)
                results = ocr.readtext(image_array, text_threshold=0.5)
                if not results:
                    text = ""
                else:
                    text_lines = [result[1] for result in results if result[1].strip()]
                    text = "\n".join(text_lines)
                pages.append({"page_num": page_num, "text": text.strip()})
            
            print(f"\n  ✓ OCR完成 {len(pages)} 页")
            return pages if any(p["text"].strip() for p in pages) else None
            
        except Exception as e:
            error_msg = str(e)[:100]
            print(f"  ⚠️ OCR解析失败: {error_msg}")
            return None
    
    def _parse_pdf(self, filepath: str) -> List[Dict]:
        """PDF解析主入口（三重降级策略）"""
        print(f"  尝试解析PDF...")
        
        pages = self._parse_pdf_with_pdfplumber(filepath)
        if pages:
            print(f"  ✓ pdfplumber成功: {len(pages)} 页 (文本型PDF)")
            return pages
        
        print(f"  → 回退到PyMuPDF...")
        pages = self._parse_pdf_with_fitz(filepath)
        if pages:
            print(f"  ✓ PyMuPDF成功: {len(pages)} 页 (文本型PDF)")
            return pages
        
        if USE_OCR:
            print(f"  → 回退到OCR（扫描件PDF）...")
            pages = self._parse_pdf_with_ocr(filepath)
            if pages:
                print(f"  ✓ OCR成功: {len(pages)} 页 (扫描件PDF)")
                return pages
        
        raise ValueError("PDF解析失败：文本型与OCR方案均未提取到有效文本")
    
    def _parse_text(self, filepath: str) -> List[Dict]:
        """文本文件解析"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(filepath, 'r', encoding='gbk') as f:
                text = f.read()
        
        page_size = 2000
        pages = []
        for i in range(0, len(text), page_size):
            pages.append({
                "page_num": i // page_size + 1,
                "text": text[i:i+page_size].strip()
            })
        return pages
    
    def _parse_docx(self, filepath: str) -> Optional[List[Dict]]:
        """Word .docx 文件解析"""
        if not HAS_DOCX:
            print("  ⚠️ 未安装python-docx，跳过.docx文件")
            return None
        
        try:
            doc = Document(filepath)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        full_text += "\n\n表格: " + " | ".join(cells)
            
            if not full_text.strip():
                print("  ⚠️ .docx文件内容为空")
                return None
            
            return [{"page_num": 1, "text": full_text.strip()}]
        except Exception as e:
            print(f"  ⚠️ .docx解析失败: {str(e)[:80]}")
            return None
    
    def _parse_doc(self, filepath: str) -> Optional[List[Dict]]:
        """Word .doc 文件解析"""
        if not HAS_TEXTRACT:
            print("  ⚠️ 未安装textract，无法处理.doc文件（建议转换为.docx）")
            return None
        
        try:
            text = textract.process(str(filepath), encoding='utf-8')
            decoded_text = text.decode('utf-8').strip()
            if not decoded_text:
                print("  ⚠️ .doc文件内容为空")
                return None
            return [{"page_num": 1, "text": decoded_text}]
        except Exception as e:
            print(f"  ⚠️ .doc解析失败: {str(e)[:80]}")
            return None
    
    def _save_extracted_json(self, filepath: Path, pages: List[Dict], business_domain: str):
        """保存提取的文本为JSON（便于人工校对）"""
        doc_id = filepath.stem
        doc_name = filepath.name
        
        # 构建JSON结构（保留所有元数据）
        output_data = {
            "metadata": {
                "doc_id": doc_id,
                "doc_name": doc_name,
                "business_domain": business_domain,
                "permission_tag": "public",
                "source_path": str(filepath.absolute()),
                "extracted_at": str(Path(filepath).stat().st_mtime)
            },
            "pages": pages  # 每页包含 {"page_num": int, "text": str}
        }
        
        # 生成安全文件名
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', doc_id)
        output_path = self.text_output_dir / f"{safe_name}.json"
        
        # 保存JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        
        print(f"  ✓ 原始文本已保存: {output_path.name} ({len(pages)} 页)")
        print(f"     💡 人工校对提示: 用文本编辑器打开此JSON，直接修改 pages[].text 内容")
    
    def _file_to_text(self, filepath: str):
        """统一入口：解析文档并保存JSON"""
        filepath = Path(filepath)
        doc_id = filepath.stem
        doc_name = filepath.name
        business_domain = self._infer_domain(doc_name)
        
        # 解析内容
        suffix = filepath.suffix.lower()
        if suffix == '.pdf':
            pages = self._parse_pdf(str(filepath))
        elif suffix in ['.txt', '.md']:
            pages = self._parse_text(str(filepath))
        elif suffix == '.docx':
            pages = self._parse_docx(str(filepath))
        elif suffix == '.doc':
            pages = self._parse_doc(str(filepath))
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")
        
        if not pages:
            raise ValueError("未提取到有效文本内容")
        
        # 保存JSON
        self._save_extracted_json(filepath, pages, business_domain)
        print(f"  → 提取完成: {len(pages)} 页 | 领域: {business_domain}")
    
    def extract_all_documents(self, doc_dir: str = DOC_DIR):
        """批量提取文档文本"""
        doc_dir = Path(doc_dir)
        if not doc_dir.exists():
            doc_dir.mkdir(parents=True, exist_ok=True)
            print(f"⚠️ 文档目录不存在，已创建: {doc_dir}")
            return
        
        # 查找支持的文件
        supported_exts = {'.pdf', '.txt', '.md', '.doc', '.docx'}
        doc_files = [
            f for f in doc_dir.rglob("*") 
            if f.is_file() 
            and f.suffix.lower() in supported_exts
            and not f.name.startswith('.')
        ]
        
        if not doc_files:
            print(f"⚠️ 未找到文档（支持格式: {', '.join(supported_exts)}）")
            print(f"   请将文档放入 {doc_dir} 及其子目录")
            return
        
        print(f"✓ 找到 {len(doc_files)} 个文档（含子目录），开始提取文本...")
        
        for filepath in sorted(doc_files):
            print(f"\n📄 处理: {filepath.relative_to(doc_dir)}")
            try:
                self._file_to_text(str(filepath))
            except Exception as e:
                print(f"  ✗ 失败: {str(e)}")
                continue
        
        print(f"\n🎉 文本提取完成！校对文件保存在: {self.text_output_dir.absolute()}")
        print(f"💡 下一步: 人工校对 extracted_text/*.json 后，运行 vectorize_text.py 生成向量库")


# ============ 主程序 ============
if __name__ == "__main__":
    # 检查文档目录
    doc_dir = Path(DOC_DIR)
    supported_exts = {'.pdf', '.txt', '.md', '.doc', '.docx'}
    has_files = any(
        f.is_file() and f.suffix.lower() in supported_exts and not f.name.startswith('.')
        for f in doc_dir.rglob("*")
    )
    
    if not has_files:
        print(f"⚠️ 警告: 文档目录 '{DOC_DIR}' 中未找到支持的文档")
        print(f"   支持格式: {', '.join(supported_exts)}")
        exit(1)
    
    # 执行提取
    extractor = TextExtractor()
    extractor.extract_all_documents()